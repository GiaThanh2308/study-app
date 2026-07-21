"""
Đọc file .docx (Word) bằng python-docx — trích text theo đoạn văn (paragraph) và bảng,
chia thành các chunk. KHÔNG có khái niệm "số trang" thật như PDF (số trang trong Word
chỉ xuất hiện lúc in/hiển thị, không cố định trong file) — nên trích dẫn dùng "Phần X"
(theo thứ tự nội dung), không phải "Trang X", để tránh gây hiểu lầm.

Giới hạn: chưa xử lý hình ảnh/đồ thị nhúng trong DOCX — nếu file DOCX có hình quan
trọng, nội dung hình đó sẽ bị bỏ qua (khác với PDF đã có xử lý hình ở Giai đoạn 4 bổ sung).
"""
from docx import Document as DocxDocument

CHUNK_SIZE = 800
CHUNK_OVERLAP = 150


def _extract_full_text(docx_path: str) -> str:
    doc = DocxDocument(docx_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = text.strip()
    if len(text) <= chunk_size:
        return [text] if text else []

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


def chunk_docx(docx_path: str, progress_callback=None) -> list[dict]:
    """
    Trả về [{"text": "...", "page": 1}, ...] — "page" ở đây thực chất là số thứ tự
    "Phần" trong file, KHÔNG phải số trang thật, giữ tên field "page" để tái dùng
    chung schema với vector_store.add_chunks (PDF), tránh phải sửa nhiều nơi khác.
    """
    full_text = _extract_full_text(docx_path)
    chunks = _chunk_text(full_text)

    total = len(chunks)
    result = []
    for i, c in enumerate(chunks):
        result.append({"text": c, "page": i + 1})
        if progress_callback:
            progress_callback(i + 1, total)

    return result
